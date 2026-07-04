#!/usr/bin/env python3
"""Convert researchwrite proposal markdown to properly formatted .docx.

Usage: python3 scripts/build_proposal_docx.py <input.md> [output.docx]

Formatting standard (JL's academic proposals):
- Title: centered, Times New Roman + 宋体, 18pt bold
- Headings: Times New Roman + 宋体, 16/14/12pt bold, black
- Body: Times New Roman + 宋体, 12pt, 1.5 line spacing, NO first-line indent
- Margins: 1 inch all sides
- Tables: 10pt, Table Grid style
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

inpath = sys.argv[1]
outpath = sys.argv[2] if len(sys.argv) > 2 else inpath.replace('.md', '.docx')

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Normal style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Cm(0)

# Heading styles
for level, size in [(1, 16), (2, 14), (3, 12)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 6)
    hs.paragraph_format.space_after = Pt(12 if level == 1 else 6 if level == 2 else 3)


def bp(text, bold=False, size=12, italic=False, color=None):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    run = p.add_run('• ' + text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)


def numbered(text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)


def table(headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()


# Parse markdown
with open(inpath, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buf = []
first_title = True

while i < len(lines):
    line = lines[i].rstrip('\n')

    if not line.strip():
        i += 1
        continue

    if line.strip().startswith('```'):
        if in_code:
            txt = '\n'.join(code_buf)
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(txt)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if first_title and line.startswith('# ') and not line.startswith('## '):
        title_text = line[2:].strip()
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        first_title = False
        i += 1
        continue

    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue

    if line.startswith('## ') and not line.startswith('### '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue

    if line.startswith('### ') and not line.startswith('#### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    if line.startswith('#### '):
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(line[5:].strip())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1
        continue

    if line.strip() == '---':
        doc.add_paragraph()
        i += 1
        continue

    if line.strip().startswith('|') and not line.strip().startswith('|---'):
        tbl_lines = [line]
        j = i + 1
        if j < len(lines) and '|---' in lines[j]:
            j += 1
        while j < len(lines) and lines[j].strip().startswith('|'):
            tbl_lines.append(lines[j].rstrip('\n'))
            j += 1
        if len(tbl_lines) >= 2:
            headers = [c.strip() for c in tbl_lines[0].split('|')[1:-1]]
            rows = [[c.strip() for c in l.split('|')[1:-1]] for l in tbl_lines[1:]]
            table(headers, rows)
        i = j
        continue

    if line.strip().startswith('**') and line.strip().endswith('**'):
        bp(line.strip()[2:-2], bold=True)
        i += 1
        continue

    if line.strip().startswith('- '):
        bullet(line.strip()[2:])
        i += 1
        continue

    if re.match(r'^\d+\.', line.strip()):
        numbered(line.strip())
        i += 1
        continue

    if line.startswith('>'):
        bp(line[1:].strip(), italic=True, size=10, color=RGBColor(80, 80, 80))
        i += 1
        continue

    bp(line.strip())
    i += 1

doc.save(outpath)
print(f"OK → {outpath}")
