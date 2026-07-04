#!/usr/bin/env python3
"""Read-only QA helper for Chinese review manuscripts in DOCX format."""

from __future__ import annotations

import re
import sys
import zipfile
from io import BytesIO
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except Exception as exc:  # pragma: no cover - environment report
    print(f"[ERROR] python-docx is required: {exc}")
    sys.exit(2)

try:
    from PIL import Image
except Exception:  # pragma: no cover - optional dependency
    Image = None


EMU_PER_INCH = 914400
REF_START_RE = re.compile(r"^\s*\[(\d+)\]")
NUM_CIT_RE = re.compile(r"\[(\d+(?:\s*[-,, ]\s*\d+)*)\]")
HEADING_RE = re.compile(r"^(English text|[1-9]\d?(?:\.\d+)*\s+\S+)")
CAPTION_RE = re.compile(r"^(English text\s*\d+|Figure\s+\d+|English text\s*\d+|Table\s+\d+)")
FORBIDDEN = [
    "English text",
    "English text",
    "English text",
    "AIEnglish text",
    "English text",
    "English text",
    "English text",
    "English text",
]


def iter_table_paragraphs(doc):
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            for cell_idx, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    yield table_idx, row_idx, cell_idx, paragraph


def has_child(element, local_name: str) -> bool:
    for child in element:
        if child.tag.endswith("}" + local_name):
            return True
    return False


def expand_citations(text: str) -> set[int]:
    refs: set[int] = set()
    for match in NUM_CIT_RE.finditer(text):
        token = match.group(1).replace(", ", ",")
        for part in token.split(","):
            part = part.strip()
            if "-" in part:
                left, right = [p.strip() for p in part.split("-", 1)]
                if left.isdigit() and right.isdigit():
                    refs.update(range(int(left), int(right) + 1))
            elif part.isdigit():
                refs.add(int(part))
    return refs


def paragraph_texts(doc):
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for _, _, _, paragraph in iter_table_paragraphs(doc):
        yield paragraph.text


def reference_entries(doc):
    entries: dict[int, list[str]] = {}
    current = None
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = REF_START_RE.match(text)
        if match:
            in_refs = True
            current = int(match.group(1))
            entries.setdefault(current, []).append(text)
        elif in_refs and current is not None and text:
            entries[current].append(text)
    return entries


def body_citations(doc):
    refs: set[int] = set()
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if REF_START_RE.match(text.strip()):
            in_refs = True
        if not in_refs:
            refs.update(expand_citations(text))
    for _, _, _, paragraph in iter_table_paragraphs(doc):
        refs.update(expand_citations(paragraph.text))
    return refs


def image_report(doc):
    reports = []
    for idx, shape in enumerate(doc.inline_shapes, start=1):
        width_in = shape.width / EMU_PER_INCH if shape.width else 0
        height_in = shape.height / EMU_PER_INCH if shape.height else 0
        pixels = None
        dpi = None
        if Image is not None:
            rid = None
            for element in shape._inline.iter():
                if element.tag.endswith("}blip"):
                    rid = element.get(qn("r:embed"))
                    break
            if rid and rid in doc.part.related_parts:
                blob = doc.part.related_parts[rid].blob
                try:
                    with Image.open(BytesIO(blob)) as img:
                        pixels = img.size
                    if width_in and height_in:
                        dpi = (pixels[0] / width_in, pixels[1] / height_in)
                except Exception:
                    pass
        reports.append((idx, pixels, width_in, height_in, dpi))
    return reports


def xml_footnotes(path: Path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return xml.count("w:footnoteReference"), "footnoteRef" in xml


def table_report(doc):
    rows = []
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.trPr
            cant_split = tr_pr is not None and has_child(tr_pr, "cantSplit")
            tbl_header = tr_pr is not None and has_child(tr_pr, "tblHeader")
            if row_idx == 1 or not cant_split:
                rows.append((table_idx, row_idx, cant_split, tbl_header))
    return rows


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: python docx_review_qa.py <manuscript.docx>")
        return 2

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"[ERROR] File not found: {path}")
        return 2

    doc = Document(str(path))
    texts = list(paragraph_texts(doc))
    refs = reference_entries(doc)
    cited = body_citations(doc)
    ref_nums = sorted(refs)

    print(f"[DOCX] {path}")
    print(f"[COUNTS] paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)} references={len(ref_nums)}")

    print("\n[HEADINGS]")
    for text in texts:
        stripped = text.strip()
        if HEADING_RE.match(stripped):
            print(f"  {stripped}")

    print("\n[CAPTIONS]")
    for text in texts:
        stripped = text.strip()
        if CAPTION_RE.match(stripped):
            print(f"  {stripped}")

    print("\n[REFERENCE CHECK]")
    if ref_nums:
        expected = list(range(ref_nums[0], ref_nums[-1] + 1))
        missing = sorted(set(expected) - set(ref_nums))
        print(f"  list range: {ref_nums[0]}-{ref_nums[-1]}")
        print(f"  missing in list range: {missing if missing else 'none'}")
    uncited = sorted(set(ref_nums) - cited)
    missing_entries = sorted(cited - set(ref_nums))
    print(f"  cited numbers: {sorted(cited)}")
    print(f"  cited but missing from reference list: {missing_entries if missing_entries else 'none'}")
    print(f"  listed but not cited in body/tables: {uncited if uncited else 'none'}")

    print("\n[FORBIDDEN / PROCESS WORDING]")
    hits = []
    for word in FORBIDDEN:
        count = sum(text.count(word) for text in texts)
        if count:
            hits.append((word, count))
    if hits:
        for word, count in hits:
            print(f"  [WARN] {word}: {count}")
    else:
        print("  none")

    print("\n[FOOTNOTE RESIDUE]")
    visible = sum("footnoteRef" in text for text in texts)
    try:
        xml_count, xml_visible = xml_footnotes(path)
        print(f"  visible footnoteRef paragraphs: {visible}")
        print(f"  document.xml w:footnoteReference count: {xml_count}; literal footnoteRef in XML: {xml_visible}")
    except Exception as exc:
        print(f"  [WARN] XML footnote scan failed: {exc}")

    print("\n[TABLE PAGINATION]")
    for table_idx, row_idx, cant_split, tbl_header in table_report(doc):
        flag = "OK" if cant_split else "WARN"
        header = "; header repeats" if row_idx == 1 and tbl_header else ""
        print(f"  [{flag}] table {table_idx} row {row_idx}: cantSplit={cant_split}{header}")

    print("\n[IMAGES]")
    for idx, pixels, width_in, height_in, dpi in image_report(doc):
        if pixels and dpi:
            print(f"  image {idx}: {pixels[0]}x{pixels[1]} px, {width_in:.2f}x{height_in:.2f} in, effective dpi {dpi[0]:.0f}x{dpi[1]:.0f}")
        else:
            print(f"  image {idx}: embedded size {width_in:.2f}x{height_in:.2f} in; pixel/DPI unavailable")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
