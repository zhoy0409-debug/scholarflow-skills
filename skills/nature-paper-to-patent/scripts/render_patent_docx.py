#!/usr/bin/env python3
"""Render a structured Chinese patent draft JSON file as DOCX."""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from math_to_omml import latex_to_omml


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 16 if level == 1 else 14, bold=True)


def add_body(document: Document, text: str, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(str(text))
    set_run_font(run, "宋体", 12)


def add_equation(document: Document, equation: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    latex = equation.get("latex")
    if not latex:
        raise ValueError(
            f"Equation {equation.get('number')} has no latex source for native Office Math"
        )
    paragraph._p.append(latex_to_omml(latex))
    number_run = paragraph.add_run(f"    ({equation.get('number', '')})")
    set_run_font(number_run, "Cambria Math", 12)
    description = equation.get("description")
    if description:
        add_body(document, description)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number(section)


def add_claims(document: Document, claims: list[dict]) -> None:
    add_heading(document, "权利要求书")
    for claim in claims:
        add_body(document, f"{claim['number']}. {claim['text']}", indent=False)


def add_specification(document: Document, data: dict, figure_dir: Path | None = None) -> None:
    add_heading(document, "说明书")
    add_heading(document, data.get("title", "[TO CONFIRM: title]"), level=2)
    spec = data.get("specification", {})
    sections = [
        ("技术领域", spec.get("technical_field", [])),
        ("背景技术", spec.get("background", [])),
    ]
    invention = spec.get("invention_content", {})
    sections.extend(
        [
            ("发明内容", invention.get("problem", []) + invention.get("solution", [])),
            ("有益效果", invention.get("beneficial_effects", [])),
            ("附图说明", spec.get("figure_descriptions", [])),
        ]
    )
    for heading, paragraphs in sections:
        add_heading(document, heading, level=2)
        for paragraph in paragraphs:
            add_body(document, paragraph)

    equations = spec.get("equations", [])
    if equations:
        add_heading(document, "公式及符号说明", level=2)
        for equation in equations:
            add_equation(document, equation)

    if figure_dir:
        figures = data.get("figures", [])
        if figures:
            add_heading(document, "说明书附图", level=2)
        for figure in figures:
            image = figure_dir / f"figure-{figure['number']}.png"
            if not image.exists():
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image), width=Cm(14))
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f"图{figure['number']} {figure.get('title', '')}")
            set_run_font(caption_run, "宋体", 11)

    add_heading(document, "具体实施方式", level=2)
    for embodiment in spec.get("embodiments", []):
        add_heading(document, embodiment.get("heading", "实施例"), level=2)
        for paragraph in embodiment.get("paragraphs", []):
            add_body(document, paragraph)


def add_figure(
    document: Document,
    figure: dict,
    figure_dir: Path | None,
    heading: str | None = None,
) -> bool:
    if not figure_dir:
        return False
    image = figure_dir / f"figure-{figure['number']}.png"
    if not image.exists():
        return False
    if heading:
        add_heading(document, heading, level=2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image), width=Cm(14))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(f"图{figure['number']} {figure.get('title', '')}")
    set_run_font(caption_run, "宋体", 11)
    return True


def abstract_figure(data: dict) -> dict | None:
    number = data.get("abstract_figure_number")
    return next(
        (figure for figure in data.get("figures", []) if figure.get("number") == number),
        None,
    )


def add_abstract(
    document: Document,
    data: dict,
    figure_dir: Path | None = None,
    include_figure: bool = True,
) -> None:
    add_heading(document, "说明书摘要")
    add_heading(document, data.get("title", "[TO CONFIRM: title]"), level=2)
    add_body(document, data.get("abstract", ""), indent=False)
    figure = abstract_figure(data)
    if include_figure and figure:
        add_figure(document, figure, figure_dir, heading="摘要附图")


def add_abstract_figure(
    document: Document,
    data: dict,
    figure_dir: Path | None,
) -> None:
    add_heading(document, "摘要附图")
    figure = abstract_figure(data)
    if figure:
        add_figure(document, figure, figure_dir)


def add_review_appendix(document: Document, data: dict) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "起草审查附录")
    metadata = data.get("metadata", {})
    for key, value in metadata.items():
        add_body(document, f"{key}: {value}", indent=False)

    add_heading(document, "前提假设", level=2)
    for item in data.get("assumptions", []):
        add_body(document, f"- {item}", indent=False)

    concept = data.get("invention_concept", {})
    add_heading(document, "发明构思", level=2)
    for key in ("technical_problem", "technical_means", "technical_effect"):
        if concept.get(key):
            add_body(document, f"{key}: {concept[key]}", indent=False)

    ledger = data.get("evidence_ledger", [])
    if ledger:
        add_heading(document, "证据台账", level=2)
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ("ID", "技术特征", "来源位置", "技术作用", "效果", "支持状态")
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        keys = ("id", "feature", "source_location", "technical_role", "effect", "support_status")
        for item in ledger:
            cells = table.add_row().cells
            for cell, key in zip(cells, keys):
                cell.text = str(item.get(key, ""))

    audit = data.get("audit", {})
    for heading, key in (
        ("支持性检查", "support_findings"),
        ("一致性检查", "consistency_findings"),
        ("发明人待确认问题", None),
    ):
        add_heading(document, heading, level=2)
        items = data.get("inventor_questions", []) if key is None else audit.get(key, [])
        for item in items:
            add_body(document, f"- {item}", indent=False)


def validate(data: dict) -> None:
    required = ("title", "claims", "specification", "abstract")
    missing = [key for key in required if key not in data]
    if missing:
        raise ValueError(f"Missing required keys: {', '.join(missing)}")
    numbers = [claim.get("number") for claim in data["claims"]]
    if numbers != list(range(1, len(numbers) + 1)):
        raise ValueError(f"Claim numbers must be consecutive integers starting at 1: {numbers}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("draft", type=Path, help="UTF-8 JSON draft")
    parser.add_argument("--output", type=Path, required=True, help="Output DOCX path")
    parser.add_argument(
        "--part",
        choices=("all", "claims", "specification", "abstract", "abstract-figure"),
        default="all",
        help="Document part to render",
    )
    parser.add_argument(
        "--figure-dir",
        type=Path,
        help="Directory containing figure-N.png images for the specification",
    )
    args = parser.parse_args()

    data = json.loads(args.draft.read_text(encoding="utf-8"))
    validate(data)
    document = Document()
    configure(document)
    if args.part == "claims":
        add_claims(document, data["claims"])
    elif args.part == "specification":
        add_specification(document, data, args.figure_dir)
    elif args.part == "abstract":
        add_abstract(document, data, args.figure_dir)
    elif args.part == "abstract-figure":
        add_abstract_figure(document, data, args.figure_dir)
    else:
        add_claims(document, data["claims"])
        document.add_page_break()
        add_specification(document, data, args.figure_dir)
        document.add_page_break()
        add_abstract(document, data, args.figure_dir)
        add_review_appendix(document, data)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
